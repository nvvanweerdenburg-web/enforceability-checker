"""Extract plain text from uploaded PDF / DOCX / TXT files."""

from __future__ import annotations

import io
from pathlib import Path


def extract_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _pdf(data)
    if suffix in {".docx"}:
        return _docx(data)
    if suffix in {".txt", ".md"}:
        return data.decode("utf-8", errors="replace")
    raise ValueError(f"Unsupported file type: {suffix}. Upload PDF, DOCX, TXT, or MD.")


def _pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    return "\n\n".join(parts).strip()


def _docx(data: bytes) -> str:
    import docx  # python-docx

    doc = docx.Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return "\n".join(parts).strip()


def chunk_document(text: str, target_chars: int = 1500, overlap_chars: int = 150) -> list[str]:
    """Split a long document into paragraph-aware chunks for per-chunk retrieval."""
    text = text.strip()
    if not text:
        return []

    # First split by blank lines into paragraphs.
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: list[str] = []
    buf: list[str] = []
    size = 0
    for p in paragraphs:
        if size + len(p) + 2 > target_chars and buf:
            chunks.append("\n\n".join(buf))
            # Keep an overlap tail.
            tail = chunks[-1][-overlap_chars:]
            buf = [tail] if tail else []
            size = len(tail)
        buf.append(p)
        size += len(p) + 2
    if buf:
        chunks.append("\n\n".join(buf))
    return chunks
